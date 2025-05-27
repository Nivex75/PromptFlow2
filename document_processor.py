import docx
import io
import logging
import PyPDF2
from typing import Optional, Dict, Any

class DocumentProcessor:
    """
    Document processor that handles multiple file formats including Word and PDF.
    Provides robust text extraction with fallback mechanisms.
    """

    SUPPORTED_FORMATS = ['docx', 'pdf', 'txt']

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text(self, uploaded_file) -> str:
        """
        Extract text from an uploaded document (Word, PDF, or text file)

        Args:
            uploaded_file: Streamlit UploadedFile object

        Returns:
            Extracted text as string

        Raises:
            ValueError: If file type is not supported or file is invalid
            Exception: For other processing errors
        """
        try:
            if uploaded_file is None:
                raise ValueError("No file was uploaded")

            # Log file details for debugging
            self.logger.info(f"Processing file: {uploaded_file.name}, type: {uploaded_file.type}")

            # Read the file content
            file_content = uploaded_file.read()
            if not file_content:
                raise ValueError("Uploaded file is empty")

            # Determine file type based on extension
            file_extension = uploaded_file.name.split('.')[-1].lower()

            # Validate file type
            if file_extension not in self.SUPPORTED_FORMATS:
                raise ValueError(
                    f"Unsupported file type: {file_extension}. "
                    f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
                )

            # Process based on file type
            extraction_methods = {
                'docx': self._extract_from_docx,
                'pdf': self._extract_from_pdf,
                'txt': self._extract_from_txt
            }

            extracted_text = extraction_methods[file_extension](file_content)

            # Validate extraction
            if not extracted_text or not extracted_text.strip():
                raise ValueError(f"No text could be extracted from the {file_extension} file")

            self.logger.info(f"Successfully extracted {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise
        finally:
            # Reset the file pointer for future reads
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from docx content"""
        try:
            # Create a BytesIO object
            doc_bytes = io.BytesIO(file_content)

            # Open document with python-docx
            doc = docx.Document(doc_bytes)

            # Extract text from paragraphs
            full_text = []

            # Add document title if exists
            if doc.core_properties.title:
                full_text.append(f"Title: {doc.core_properties.title}\n")

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text.append(paragraph.text)

            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                table_text = [f"\n[Table {table_idx + 1}]"]
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if len(table_text) > 1:  # Only add if table has content
                    full_text.extend(table_text)

            return "\n".join(full_text)

        except Exception as e:
            self.logger.error(f"Error extracting from DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from Word document: {str(e)}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF content using PyPDF2 with fallback options"""
        try:
            # Create a BytesIO object from file content
            pdf_bytes = io.BytesIO(file_content)

            # Try PyPDF2 first
            try:
                return self._extract_with_pypdf2(pdf_bytes)
            except Exception as e:
                self.logger.warning(f"PyPDF2 extraction failed: {str(e)}")

                # Reset stream position
                pdf_bytes.seek(0)

                # Try alternative method if available
                try:
                    import pdfplumber
                    return self._extract_with_pdfplumber(pdf_bytes)
                except ImportError:
                    self.logger.info("pdfplumber not available, using PyPDF2 result")
                except Exception as e:
                    self.logger.warning(f"pdfplumber extraction also failed: {str(e)}")

                # Return whatever we got from PyPDF2
                pdf_bytes.seek(0)
                return self._extract_with_pypdf2(pdf_bytes, strict=False)

        except Exception as e:
            self.logger.error(f"Error extracting from PDF: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_with_pypdf2(self, pdf_bytes: io.BytesIO, strict: bool = True) -> str:
        """Extract text using PyPDF2"""
        pdf_reader = PyPDF2.PdfReader(pdf_bytes, strict=strict)

        # Get document info
        full_text = []
        if pdf_reader.metadata:
            if pdf_reader.metadata.title:
                full_text.append(f"Title: {pdf_reader.metadata.title}\n")

        # Extract text from all pages
        total_pages = len(pdf_reader.pages)
        for page_num in range(total_pages):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                if page_text and page_text.strip():
                    # Add page marker for multi-page documents
                    if total_pages > 1:
                        full_text.append(f"\n--- Page {page_num + 1} ---\n")
                    full_text.append(page_text)
            except Exception as e:
                self.logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                if strict:
                    raise

        # Check if we got any text
        if not full_text or (len(full_text) == 1 and full_text[0].startswith("Title:")):
            return "The PDF appears to contain no extractable text. It may be a scanned document or image-based PDF."

        return "\n".join(full_text)

    def _extract_with_pdfplumber(self, pdf_bytes: io.BytesIO) -> str:
        """Extract text using pdfplumber (better for complex PDFs)"""
        import pdfplumber

        full_text = []
        with pdfplumber.open(pdf_bytes) as pdf:
            # Add metadata if available
            if pdf.metadata.get('Title'):
                full_text.append(f"Title: {pdf.metadata['Title']}\n")

            # Extract from each page
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    if len(pdf.pages) > 1:
                        full_text.append(f"\n--- Page {i + 1} ---\n")
                    full_text.append(page_text)

        return "\n".join(full_text) if full_text else "No text could be extracted from PDF"

    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace')

        except Exception as e:
            self.logger.error(f"Error extracting from text file: {str(e)}")
            raise ValueError(f"Failed to read text file: {str(e)}")

    def get_document_info(self, uploaded_file) -> Dict[str, Any]:
        """
        Get document information without extracting all text
        Useful for quick document preview
        """
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_content = uploaded_file.read()
            uploaded_file.seek(0)

            info = {
                'filename': uploaded_file.name,
                'size': len(file_content),
                'type': file_extension,
                'pages': 1  # Default
            }

            if file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                info['pages'] = len(pdf_reader.pages)
                if pdf_reader.metadata:
                    info['title'] = pdf_reader.metadata.get('/Title', '')
                    info['author'] = pdf_reader.metadata.get('/Author', '')

            elif file_extension == 'docx':
                doc = docx.Document(io.BytesIO(file_content))
                info['title'] = doc.core_properties.title or ''
                info['author'] = doc.core_properties.author or ''
                # Rough page estimate
                info['pages'] = max(1, len(doc.paragraphs) // 25)

            return info

        except Exception as e:
            self.logger.error(f"Error getting document info: {str(e)}")
            return {'error': str(e)}